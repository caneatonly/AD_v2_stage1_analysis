from __future__ import annotations

import copy
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
LATEX_DIR = ROOT / "paper" / "latex"
MAIN_TEX = LATEX_DIR / "main.tex"
HIGHLIGHTS_TEX = LATEX_DIR / "highlights_items.tex"
TMP_DIR = ROOT / "tmp" / "docs"
OUTPUT_DIR = ROOT / "output" / "doc"
TMP_CONVERTED_MD = TMP_DIR / "main_from_latex.md"
TMP_WORD_MD = TMP_DIR / "word_manuscript.md"
OUTPUT_DOCX = OUTPUT_DIR / "ocean_engineering_word_mirror.docx"

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def run(command: list[str], cwd: Path | None = None) -> None:
    subprocess.run(command, cwd=cwd, check=True)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def require_match(pattern: str, text: str, label: str, flags: int = 0) -> str:
    match = re.search(pattern, text, flags)
    if not match:
        raise ValueError(f"Could not extract {label} from {MAIN_TEX}")
    return match.group(1).strip()


def extract_metadata(main_tex: str) -> dict[str, object]:
    title = require_match(r"\\title\{(.+?)\}", main_tex, "title", re.S)
    abstract = require_match(
        r"\\begin\{abstract\}\s*(.+?)\s*\\end\{abstract\}",
        main_tex,
        "abstract",
        re.S,
    )
    journal = require_match(r"\\journal\{(.+?)\}", main_tex, "journal")
    keyword_block = require_match(
        r"\\begin\{keyword\}\s*(.+?)\s*\\end\{keyword\}",
        main_tex,
        "keywords",
        re.S,
    )
    authors_raw: list[str] = []
    for line in main_tex.splitlines():
        match = re.search(r"\\author(?:\[[^\]]+\])?\{(.*)\}\s*$", line.strip())
        if match:
            authors_raw.append(match.group(1))
    if not authors_raw:
        raise ValueError(f"Could not extract authors from {MAIN_TEX}")
    authors = [strip_latex_commands(item) for item in authors_raw]

    email_match = re.search(r"\\ead\{(.+?)\}", main_tex)
    email = email_match.group(1).strip() if email_match else ""

    org_match = re.search(r"organization=\{(.*?)\}", main_tex, re.S)
    city_match = re.search(r"city=\{(.*?)\}", main_tex, re.S)
    country_match = re.search(r"country=\{(.*?)\}", main_tex, re.S)
    affiliation_parts = []
    for value in (org_match, city_match, country_match):
        if value and value.group(1).strip():
            affiliation_parts.append(value.group(1).strip())
    affiliation = ", ".join(affiliation_parts)

    keywords = [item.strip() for item in re.split(r"\\sep", keyword_block) if item.strip()]
    highlights = []
    if HIGHLIGHTS_TEX.exists():
        for line in read_text(HIGHLIGHTS_TEX).splitlines():
            line = line.strip()
            if line.startswith(r"\item"):
                highlights.append(line.replace(r"\item", "", 1).strip())

    return {
        "title": latex_to_plain(title),
        "abstract": latex_to_plain(abstract),
        "journal": latex_to_plain(journal),
        "authors": [latex_to_plain(author) for author in authors],
        "email": latex_to_plain(email),
        "affiliation": latex_to_plain(affiliation),
        "keywords": [latex_to_plain(item) for item in keywords],
        "highlights": [latex_to_plain(item) for item in highlights],
    }


def latex_to_plain(text: str) -> str:
    replacements = {
        "~": " ",
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"``": '"',
        r"''": '"',
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.strip()


def strip_latex_commands(text: str) -> str:
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\[A-Za-z@]+(?:\[[^\]]*\])?\{[^{}]*\}", "", text)
    text = re.sub(r"\\[A-Za-z@]+", "", text)
    text = text.replace("{", "").replace("}", "")
    return latex_to_plain(text)


def build_converted_markdown() -> str:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    run(
        [
            "pandoc",
            "main.tex",
            "--from=latex",
            "--to=markdown",
            "--citeproc",
            "--bibliography=references.bib",
            "-o",
            str(TMP_CONVERTED_MD),
        ],
        cwd=LATEX_DIR,
    )
    converted = read_text(TMP_CONVERTED_MD)
    intro_match = re.search(r"(?m)^# Introduction\b", converted)
    if not intro_match:
        raise ValueError("Pandoc output did not contain the Introduction heading.")
    body = converted[intro_match.start() :].strip()

    refs_match = re.search(r"(?m)^:{3,}\s*\{#refs\b.*$", body)
    if refs_match:
        pre_refs = body[: refs_match.start()].rstrip()
        refs = body[refs_match.start() :].strip()
        body = f"{pre_refs}\n\n# References {{.unnumbered}}\n\n{refs}\n"

    return body


def yaml_escape(text: str) -> str:
    return text.replace('"', '\\"')


def assemble_word_markdown(meta: dict[str, object], body: str) -> str:
    authors = meta["authors"]
    highlights = meta["highlights"]
    keywords = meta["keywords"]

    yaml_lines = ["---", f'title: "{yaml_escape(meta["title"])}"', "author:"]
    for author in authors:
        yaml_lines.append(f'  - "{yaml_escape(author)}"')
    yaml_lines.extend(["...", ""])

    frontmatter = [
        f"*{meta['journal']}*",
        "",
    ]
    if meta["affiliation"]:
        frontmatter.append(meta["affiliation"])
        frontmatter.append("")
    if meta["email"]:
        frontmatter.append(f"Corresponding author: {meta['email']}")
        frontmatter.append("")

    frontmatter.extend(
        [
            "# Abstract {.unnumbered}",
            "",
            str(meta["abstract"]).strip(),
            "",
            f"**Keywords:** {'; '.join(keywords)}",
        ]
    )

    if highlights:
        frontmatter.extend(
            [
                "",
                "# Highlights {.unnumbered}",
                "",
            ]
        )
        frontmatter.extend(f"- {item}" for item in highlights)

    return "\n".join(yaml_lines + frontmatter) + "\n\n" + body.strip() + "\n"


def set_font(style, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), "Times New Roman")


def style_by_name(doc: Document, name: str):
    for style in doc.styles:
        if style.name == name:
            return style
    return None


def apply_styles(doc: Document) -> None:
    paragraph_styles = {
        "Normal": (10.0, False, False),
        "Body Text": (10.0, False, False),
        "First Paragraph": (10.0, False, False),
        "Title": (14.0, True, False),
        "Author": (10.5, False, False),
        "Subtitle": (9.5, False, True),
        "Heading 1": (11.0, True, False),
        "Heading 2": (10.0, True, False),
        "Heading 3": (9.5, True, False),
        "Bibliography": (8.5, False, False),
        "Compact": (8.5, False, False),
        "Abstract": (9.5, False, False),
        "Abstract Title": (10.0, True, False),
        "Image Caption": (8.5, False, False),
        "Table Caption": (8.5, False, False),
        "Caption": (8.5, False, False),
        "Table": (8.5, False, False),
    }

    for name, (size, bold, italic) in paragraph_styles.items():
        style = style_by_name(doc, name)
        if style is not None and style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            set_font(style, size, bold=bold, italic=italic)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
        elif paragraph.style.name in {"Author", "Subtitle"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
        elif paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
        elif paragraph.style.name in {"Bibliography", "Image Caption", "Table Caption", "Caption"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(-0.63) if paragraph.style.name == "Bibliography" else None
            paragraph.paragraph_format.left_indent = Cm(0.63) if paragraph.style.name == "Bibliography" else None
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.startswith("Corresponding author:") else WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.first_line_indent = Cm(0)

    if doc.paragraphs:
        frontmatter_boundary = None
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == "Abstract":
                frontmatter_boundary = paragraph
                break
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)


def configure_page(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def remove_children(element, tag: str) -> None:
    for child in list(element):
        if child.tag == tag:
            element.remove(child)


def set_columns(sect_pr, count: int, space_twips: int = 425) -> None:
    remove_children(sect_pr, qn("w:cols"))
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    if count > 1:
        cols.set(qn("w:space"), str(space_twips))
    sect_pr.append(cols)


def set_section_type(sect_pr, value: str) -> None:
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), value)


def insert_body_section_break(doc: Document) -> None:
    intro_paragraph = None
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text in {"1 Introduction", "Introduction"}:
            intro_paragraph = paragraph
            break
    if intro_paragraph is None:
        return

    body = doc._element.body
    body_sect_pr = body.sectPr
    if body_sect_pr is None:
        return
    set_columns(body_sect_pr, 2, 425)

    break_paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    break_paragraph.append(p_pr)

    break_sect_pr = copy.deepcopy(body_sect_pr)
    set_columns(break_sect_pr, 1)
    set_section_type(break_sect_pr, "continuous")
    p_pr.append(break_sect_pr)

    intro_paragraph._p.addprevious(break_paragraph)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0


def build_docx() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    meta = extract_metadata(read_text(MAIN_TEX))
    body = build_converted_markdown()
    word_md = assemble_word_markdown(meta, body)
    TMP_WORD_MD.write_text(word_md, encoding="utf-8")

    run(
        [
            "pandoc",
            str(TMP_WORD_MD),
            "--from=markdown",
            "--to=docx",
            "--number-sections",
            "-o",
            str(OUTPUT_DOCX),
        ],
        cwd=ROOT,
    )

    doc = Document(OUTPUT_DOCX)
    for section in doc.sections:
        configure_page(section)

    apply_styles(doc)
    insert_body_section_break(doc)
    format_tables(doc)
    doc.core_properties.title = str(meta["title"])
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_docx()
    print(output)
